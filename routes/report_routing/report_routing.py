from flask import Blueprint, render_template, request, jsonify
import os
from config import add_logger, DETAILS_PATH
from flask import send_file
from docx import Document
from io import BytesIO
import numpy as np
# from docx.shared import Inches
# from openpyxl import load_workbook
import json
import datetime
# from dotenv import load_dotenv, set_key
from flask_cors import CORS

# load_dotenv()


script_name = os.path.splitext(os.path.basename(__file__))[0]
logger = add_logger(f'logger_{script_name}', script_name)


report = Blueprint("reports", __name__)
CORS(report)


AVERAGE_COLUMN = 'Average'
ALLOWANCE_COLUMN = 'Allowance'
GEOMETRY_COLUMN = 'Geometry'


@report.route('/ticket_form')
def ticket_form():
    return render_template('ticket.html')


@report.route('/get_details')
def get_details():
    with open(DETAILS_PATH, 'r', encoding='utf-8') as file:
        details = json.load(file)
        print(details)
    return jsonify(details)


@report.route('/select_detail', methods=['POST', 'GET'])
def select_detail():
    detail = request.json.get('detail')
    print(f"DETAIL_SD: {detail}")
    with open(DETAILS_PATH, 'r', encoding='utf-8') as file:
        details = json.load(file)
    response = details["details_map"].get(detail)
    print(f"RESPONCE_SD:{response}")
    return jsonify(response), 200


def create_table(data_json):
    def calculate_average(numbers):
        # Преобразование всех элементов в числа, игнорируя нечисловые значения
        filtered_numbers = []
        for x in numbers:
            try:
                filtered_numbers.append(float(x))
            except ValueError:
                continue
        # Вычисление среднего значения, если есть числа
        return round(np.mean(filtered_numbers), 2) if filtered_numbers else None

    for title, section in data_json.items():
        for item in section.values():
            header = item[0]
            header.extend(['Average', 'Allowance', 'Geometry'])
            for row in item[1:]:
                # Преобразование строковых измерений в числа
                measurements = [x.replace(',', '.') for x in row[2]]
                average = calculate_average(measurements)
                row.extend([average, '', ''])

    return data_json


@report.route('/put_data', methods=['POST'])
def save_table():
    try:
        data = request.get_json()
        print(data)
        if not data:
            return jsonify({'error': 'Отсутствуют данные в запросе'}), 400
        table_name = next(iter(data))
        create_table(data)
        today_date = datetime.date.today().strftime("%d-%m-%Y")
        directory = f'data/json/{today_date}/'
        if not os.path.exists(directory):
            os.makedirs(directory)

        file_path = f'{directory}/{table_name}.json'
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        print("Save data")
        return jsonify({'message': 'Данные успешно сохранены'}), 200

    except Exception as e:
        print(f"Ошибка при сохранении данных: {str(e)}")
        return jsonify({'error': 'Внутренняя ошибка сервера'}), 500


@report.route('/custom_table', methods=['POST'])
def save_custom_table():
    try:
        data = request.get_json()
        print(data)
        if not data:
            return jsonify({'error': 'Отсутствуют данные в запросе'}), 400
        table_name = next(iter(data))
        print(f"TABLE_NAME: {table_name}")
        detail_name = next(iter(data[table_name]))
        print(f"DETAIL_NAME: {detail_name}")
        is_checked = data[table_name].get('isChecked')
        if is_checked:
            rows = int(data[table_name]['rows'])
            print(f"ROWS: {rows}")
            new_detail = {'col': 3, 'row': rows}
            new_name = {"name": detail_name, "value": detail_name}
            with open(DETAILS_PATH, 'r', encoding='utf-8') as file:
                details = json.load(file)
            print(f"DETAIL_CT: {details}")
            details['details_map'][detail_name] = new_detail
            details['detail_names'].append(new_name)
            with open(DETAILS_PATH, 'w', encoding='utf-8') as file:
                json.dump(details, file, ensure_ascii=False, indent=4)

        data[table_name] = {key: value for key, value in data[table_name].items() if key not in ['isChecked', 'rows']}
        table = create_table(data)

        today_date = datetime.date.today().strftime("%d-%m-%Y")
        directory = f'data/json/{today_date}/'
        if not os.path.exists(directory):
            os.makedirs(directory)

        file_path = f'{directory}/{table_name}.json'
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(table, f, ensure_ascii=False, indent=4)

        return jsonify({'message': 'Данные успешно сохранены'}), 200

    except Exception as e:
        print(e)
        print(f"Ошибка при сохранении данных: {str(e)}")
        return jsonify({'error': 'Внутренняя ошибка сервера'}), 500


# @report.route('/submit_ticket', methods=['POST'])
# def submit_ticket():
#     title = request.form['ticket_title']
#     description = request.form['ticket_description']
#     num_rows = int(request.form['table_rows'])
#     num_cols = int(request.form['table_cols'])

#     return f'<h1>Ticket submitted:</h1><p>Title: {title}</p><p>Description: {description}</p><p>Table: {num_rows}x{num_cols}</p>'


@report.route('/')
def reports():
    return render_template('reports.html')


@report.route('/custom_template')
def custom_template():
    return render_template('custom_template.html')


@report.route('/export_docx')
def fill_document():
    report_file = BytesIO()
    doc = Document("Template.docx")

    data = {
        'date': '21 марта 2024',
        'names': 'Налётов В.А, Петров П.П.',
        'name_machine': 'Оборудование X',
        'company': 'Компания ABC',
        'location': 'ул. Пушкина, д.10',
        'start_time': '10:00',
        'end_time': '12:00',
        'temp': '25',
    }

    data_pipe = {
        'Труба 1': {
            'nominal1': [356],
            'Measurements1': list(range(300, 311)),
            'Average1': [433],
            'Allowance1': ["что-то"],
            'Geometry1': ["что-то"]
        }
    }

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if '{{' + key + '}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{' + key + '}}', value)
        for key, value in data_pipe.items():
            if '{{' + key + '}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{' + key + '}}', value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for name, pipe_data in data_pipe.items():
                    for key, value in pipe_data.items():
                        placeholder = '{{' + key + '}}'
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, ', '.join(map(str, value)))

    doc.save(report_file)
    report_file.seek(0)

    return send_file(report_file, as_attachment=True, download_name='report.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
